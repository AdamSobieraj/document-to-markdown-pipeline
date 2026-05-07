import io
import logging
import os
from typing import List, Union
import docx
from langchain_core.documents import Document
from .base_parser import BaseDocumentParser

logger = logging.getLogger(__name__)


class DocxParser(BaseDocumentParser):
    """
    Parser dla dokumentów Word, konwertuje style na formatowanie Markdown.
    Odczytuje fizyczne i logiczne podziały stron.
    """

    def parse(self, file_source: Union[str, io.BytesIO, bytes], **kwargs) -> List[Document]:
        documents = []
        source_name = os.path.basename(file_source) if isinstance(file_source, str) else "strumień pamięci S3"

        try:
            # Upewniamy się, że surowe bajty z S3 są czytane jako strumień pliku
            if isinstance(file_source, bytes):
                file_source = io.BytesIO(file_source)

            # python-docx potrafi czytać strumienie z pamięci RAM (BytesIO) lub ścieżki (str)
            doc = docx.Document(file_source)

            current_page_text = []
            current_char_count = 0
            page_num = 1
            MAX_CHARS_PER_PAGE = 2500

            for para in doc.paragraphs:
                # python-docx automatycznie mapuje prefiks 'w:', nie używamy argumentu namespaces!
                hard_breaks = para._element.xpath('.//w:br[@w:type="page"]')
                rendered_breaks = para._element.xpath('.//w:lastRenderedPageBreak')

                # Jeśli trafimy na znacznik strony lub przekroczymy limit znaków na stronę
                if hard_breaks or rendered_breaks or current_char_count >= MAX_CHARS_PER_PAGE:
                    joined_text = "\n\n".join(current_page_text).strip()
                    if joined_text:
                        documents.append(Document(page_content=joined_text, metadata={"page_number": page_num}))

                    # Reset na nową stronę
                    current_page_text = []
                    current_char_count = 0
                    page_num += 1

                text = para.text.strip()
                if not text:
                    continue

                # --- Konwersja na formatowanie Markdown ---
                style_name = para.style.name.lower()
                md_text = text

                if style_name.startswith('heading'):
                    level = style_name.replace('heading', '').strip()
                    if level.isdigit():
                        md_text = f"{'#' * int(level)} {text}"
                elif 'list paragraph' in style_name or 'bullet' in style_name:
                    md_text = f"- {text}"
                elif 'num' in style_name:
                    md_text = f"1. {text}"

                current_page_text.append(md_text)
                current_char_count += len(md_text)

            # Zapisz resztkę tekstu jako ostatnią stronę (jeśli plik nie kończył się podziałem)
            if current_page_text:
                joined_text = "\n\n".join(current_page_text).strip()
                if joined_text:
                    documents.append(Document(page_content=joined_text, metadata={"page_number": page_num}))

            return documents

        except Exception as e:
            logger.error(f"Błąd parsowania DOCX ({source_name}): {e}")
            return []